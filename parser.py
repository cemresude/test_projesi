import PyPDF2
from docx import Document
import sys
import os


def extract_text_from_pdf(pdf_path):
    """
    PDF dosyasından metin çıkarır.
    
    Args:
        pdf_path: PDF dosyasının yolu
    
    Returns:
        Çıkarılan metin
    """
    extracted_text = ""
    
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(pdf_reader.pages)
        
        print(f"PDF'de {num_pages} sayfa bulundu.")
        
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text:
                extracted_text += f"\n--- Sayfa {page_num + 1} ---\n"
                extracted_text += text
    
    return extracted_text


def extract_text_from_docx(docx_path):
    """
    DOCX dosyasından metin çıkarır.
    
    Args:
        docx_path: DOCX dosyasının yolu
    
    Returns:
        Çıkarılan metin
    """
    doc = Document(docx_path)
    extracted_text = ""
    
    # Paragrafları çıkar
    for para in doc.paragraphs:
        if para.text.strip():
            extracted_text += para.text + "\n"
    
    # Tablolardaki metinleri de çıkar
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                extracted_text += " | ".join(row_text) + "\n"
    
    print(f"DOCX dosyasından metin çıkarıldı.")
    return extracted_text


def extract_text(file_path, output_path=None):
    """
    PDF veya DOCX dosyasından metin çıkarır ve TXT olarak kaydeder.
    
    Args:
        file_path: Dosyanın yolu (PDF veya DOCX)
        output_path: Çıktı TXT dosyasının yolu (opsiyonel)
    
    Returns:
        Çıkarılan metin
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")
    
    # Dosya uzantısını kontrol et
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.pdf':
        extracted_text = extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        if ext == '.doc':
            print("Uyarı: .doc dosyaları tam desteklenmeyebilir. .docx formatını tercih edin.")
        extracted_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Desteklenmeyen dosya formatı: {ext}. Desteklenen formatlar: .pdf, .docx")
    
    # Çıktı dosya adını belirle
    if output_path is None:
        output_path = os.path.splitext(file_path)[0] + ".txt"
    
    # TXT olarak kaydet
    with open(output_path, 'w', encoding='utf-8') as txt_file:
        txt_file.write(extracted_text)
    
    print(f"Metin başarıyla çıkarıldı ve kaydedildi: {output_path}")
    return extracted_text


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Kullanım: python parser.py <dosya> [cikti_dosyasi.txt]")
        print("Desteklenen formatlar: .pdf, .docx")
        print("Örnek: python parser.py belge.pdf")
        print("Örnek: python parser.py rapor.docx")
        sys.exit(1)
    
    file_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    extract_text(file_path, output_path)
